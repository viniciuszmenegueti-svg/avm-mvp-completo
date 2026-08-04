"""Extract searchable text from PDF and DOCX reference documents.

The output is intended for local compliance analysis only and should not be
committed. PDF pages and DOCX paragraphs/tables retain explicit boundaries so
requirements can be traced back to their source location.
"""

from __future__ import annotations

import argparse
from pathlib import Path

import pdfplumber
from docx import Document


def extract_pdf(path: Path) -> str:
    chunks: list[str] = []
    with pdfplumber.open(path) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            text = page.extract_text(x_tolerance=2, y_tolerance=3) or ""
            chunks.append(f"\n\n===== PAGE {page_number} =====\n{text.strip()}")
    return "".join(chunks).lstrip()


def extract_docx(path: Path) -> str:
    document = Document(path)
    chunks: list[str] = []

    for index, paragraph in enumerate(document.paragraphs, start=1):
        text = paragraph.text.strip()
        if text:
            chunks.append(f"[P{index}] {text}")

    for table_index, table in enumerate(document.tables, start=1):
        chunks.append(f"\n===== TABLE {table_index} =====")
        for row_index, row in enumerate(table.rows, start=1):
            cells = [" ".join(cell.text.split()) for cell in row.cells]
            chunks.append(f"[R{row_index}] " + " | ".join(cells))

    return "\n".join(chunks)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("source_dir", type=Path)
    parser.add_argument("output_dir", type=Path)
    parser.add_argument("names", nargs="+")
    args = parser.parse_args()

    args.output_dir.mkdir(parents=True, exist_ok=True)
    for name in args.names:
        source = args.source_dir / name
        suffix = source.suffix.lower()
        if suffix == ".pdf":
            text = extract_pdf(source)
        elif suffix == ".docx":
            text = extract_docx(source)
        else:
            raise ValueError(f"Unsupported reference format: {source}")

        destination = args.output_dir / f"{source.stem}.txt"
        destination.write_text(text, encoding="utf-8")
        print(f"{source.name}: {len(text):,} chars -> {destination}")


if __name__ == "__main__":
    main()
